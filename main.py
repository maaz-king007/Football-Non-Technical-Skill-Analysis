from utils import read_video, save_video
from trackers import Tracker
import cv2
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from math import pi
from team_assigner import TeamAssigner
from player_ball_assigner import PlayerBallAssigner
from camera_movement_estimator import CameraMovementEstimator
from view_transformer import ViewTransformer
from speed_and_distance_estimator import SpeedAndDistance_Estimator
from behavior_analysis import BehaviorAnalyzer 
import copy

# =================================================================
#  HELPER: BLOCKING LOGIC
# =================================================================
def is_pass_blocked(start_point, end_point, opponent_centers, threshold=40):
    """
    Checks if a pass is blocked by an opponent.
    Threshold = distance (in pixels) from the pass line to count as a block.
    """
    p1 = np.array(start_point)
    p2 = np.array(end_point)
    line_vec = p2 - p1
    line_len = np.linalg.norm(line_vec)
    
    if line_len == 0: return False, None
    line_unit_vec = line_vec / line_len

    for opp_center in opponent_centers:
        p3 = np.array(opp_center)
        p1_to_p3 = p3 - p1
        
        # Project opponent position onto the passing line
        projection_dist = np.dot(p1_to_p3, line_unit_vec)
        
        # Only consider opponents BETWEEN the ball and the receiver
        if 0 < projection_dist < line_len:
            closest_point_on_line = p1 + (projection_dist * line_unit_vec)
            dist_to_line = np.linalg.norm(p3 - closest_point_on_line)
            
            if dist_to_line < threshold:
                return True, opp_center 

    return False, None

# =================================================================
#  HELPER: AR OVERLAY (SHADOWS & LINES)
# =================================================================
def draw_ar_overlay(frame, tracks_frame):
    # --- VISUAL SETTINGS ---
    MAX_PASS_DIST = 450           # Max distance to draw lines
    SHADOW_LENGTH = 400           # How long the grey shadows extend
    SHADOW_WIDTH_START = 15       # Width of shadow at defender's feet
    SHADOW_WIDTH_END = 80         # Width of shadow at the end (Cone shape)
    
    COLOR_OPEN = (0, 255, 255)    # Yellow (Safe Pass)
    COLOR_BLOCKED = (0, 0, 255)   # Red (Blocked Pass)
    SHADOW_COLOR = (10, 10, 10)   # Almost Black (Grey when transparent)
    SHADOW_ALPHA = 0.5            # 50% Transparency
    
    # Create overlay layer for transparency
    overlay = frame.copy()
    
    # 1. Identify Ball Carrier & Player Locations
    ball_carrier_id = None
    ball_carrier_center = None
    ball_carrier_team = None
    
    player_centers = {}
    opponent_centers = [] 

    # Collect all positions
    for track_id, data in tracks_frame.items():
        bbox = data['bbox']
        # Position = Bottom Center (Feet)
        center_x = int((bbox[0] + bbox[2]) / 2)
        feet_y = int(bbox[3]) 
        pos = (center_x, feet_y)
        
        player_centers[track_id] = pos
        
        if data.get('has_ball', False):
            ball_carrier_id = track_id
            ball_carrier_team = data.get('team')
            ball_carrier_center = pos

    # Identify Opponents (relative to ball carrier)
    if ball_carrier_team is not None:
        for track_id, data in tracks_frame.items():
            if data.get('team') != ball_carrier_team and track_id != ball_carrier_id:
                if track_id in player_centers:
                    opponent_centers.append(player_centers[track_id])

    # 2. Draw Elements if we have a Ball Carrier
    if ball_carrier_id is not None and ball_carrier_center is not None:
        
        # --- A. DRAW SHADOW ZONES (Grey Cones) ---
        for opp_center in opponent_centers:
            # Vector from Ball -> Opponent
            vec_ball_to_opp = np.array(opp_center) - np.array(ball_carrier_center)
            dist_to_opp = np.linalg.norm(vec_ball_to_opp)
            
            # Only draw shadows for opponents who are relevant (closer than 300px)
            if dist_to_opp < 300 and dist_to_opp > 0:
                unit_vec = vec_ball_to_opp / dist_to_opp
                
                # Perpendicular vector (for width)
                perp_vec = np.array([-unit_vec[1], unit_vec[0]])
                
                # Calculate the 4 points of the Shadow Trapezoid
                # Start points (at opponent)
                p1 = np.array(opp_center) + (perp_vec * SHADOW_WIDTH_START)
                p2 = np.array(opp_center) - (perp_vec * SHADOW_WIDTH_START)
                
                # End points (far away)
                end_center = np.array(opp_center) + (unit_vec * SHADOW_LENGTH)
                p3 = end_center - (perp_vec * SHADOW_WIDTH_END)
                p4 = end_center + (perp_vec * SHADOW_WIDTH_END)
                
                # Order: Top-Left -> Top-Right -> Bottom-Right -> Bottom-Left
                shadow_poly = np.array([p1, p2, p3, p4], np.int32)
                
                # Draw the filled grey polygon on the overlay
                cv2.fillPoly(overlay, [shadow_poly], SHADOW_COLOR)

        # --- B. DRAW PASSING LINES ---
        for track_id, center in player_centers.items():
            if track_id == ball_carrier_id: continue
            
            dist_px = np.linalg.norm(np.array(ball_carrier_center) - np.array(center))
            player_team = tracks_frame[track_id].get('team')
            
            # Draw line ONLY to teammates within range
            if dist_px < MAX_PASS_DIST and player_team == ball_carrier_team:
                
                # Check if pass intersects with any opponent
                blocked, _ = is_pass_blocked(ball_carrier_center, center, opponent_centers)
                
                line_color = COLOR_BLOCKED if blocked else COLOR_OPEN
                line_thickness = 2 if blocked else 3
                
                # Draw Line
                cv2.line(frame, ball_carrier_center, center, line_color, line_thickness, cv2.LINE_AA)
                # Draw Dot at feet
                cv2.circle(frame, center, 5, line_color, -1)

        # --- C. HIGHLIGHT BALL CARRIER ---
        cv2.ellipse(frame, center=ball_carrier_center, axes=(30, 15), angle=0, startAngle=0, endAngle=360, color=(0, 255, 0), thickness=2)

    # 3. APPLY TRANSPARENCY
    # Mix the overlay (with shadows) and the original frame
    cv2.addWeighted(overlay, SHADOW_ALPHA, frame, 1 - SHADOW_ALPHA, 0, frame)
    
    return frame

# =================================================================
#  PLOTTING & HELPER FUNCTIONS
# =================================================================
def smooth_speed_data(tracks, window_size=7):
    print(f"Smoothing speed data (Window Size: {window_size})...")
    for object_type in ['players']: 
        if object_type not in tracks: continue
        player_data_series = {}
        for frame_num, frame_data in enumerate(tracks[object_type]):
            for track_id, track_info in frame_data.items():
                if track_id not in player_data_series:
                    player_data_series[track_id] = {'speeds': [], 'frames': []}
                s = track_info.get('speed', 0)
                player_data_series[track_id]['speeds'].append(s)
                player_data_series[track_id]['frames'].append(frame_num)
        
        for track_id, data in player_data_series.items():
            if not data['speeds']: continue
            speed_series = pd.Series(data['speeds'])
            smoothed_speeds = speed_series.rolling(window=window_size, min_periods=1, center=False).mean()
            for i, smoothed_val in enumerate(smoothed_speeds):
                frame_idx = data['frames'][i]
                tracks[object_type][frame_idx][track_id]['speed'] = smoothed_val

def plot_speed_subplots(tracks, selected_player_ids, match_stats):
    print(f"Generating Speed Subplots for players: {selected_player_ids}...")
    player_max_speeds = []
    for pid in selected_player_ids:
        p_stats = next((item for item in match_stats if item["Player ID"] == pid), None)
        max_s = p_stats["Max Speed (km/h)"] if p_stats else 0
        player_max_speeds.append((pid, max_s))
    
    player_max_speeds.sort(key=lambda x: x[1], reverse=True)
    sorted_ids = [p[0] for p in player_max_speeds]
    
    num_players = len(sorted_ids)
    if num_players == 0: return

    fig, axes = plt.subplots(num_players, 1, figsize=(10, 3 * num_players), sharex=True)
    if num_players == 1: axes = [axes]
    colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b']
    
    for i, player_id in enumerate(sorted_ids):
        ax = axes[i]
        speeds = []
        timestamps = []
        for frame_num, frame_data in enumerate(tracks['players']):
            if player_id in frame_data:
                speeds.append(frame_data[player_id].get('speed', 0))
                timestamps.append(frame_num / 24.0) 
        
        max_speed_val = player_max_speeds[i][1]
        color = colors[i % len(colors)]
        ax.plot(timestamps, speeds, color=color, linewidth=2)
        ax.fill_between(timestamps, speeds, color=color, alpha=0.1) 
        ax.set_ylabel(f"ID #{player_id}\n(Speed km/h)", fontsize=10, rotation=0, labelpad=40, ha='center', weight='bold')
        ax.set_ylim(0, 35) 
        ax.grid(True, alpha=0.3)
        ax.text(0.02, 0.85, f"Max: {max_speed_val} km/h", transform=ax.transAxes, 
                fontsize=10, fontweight='bold', color=color, bbox=dict(facecolor='white', alpha=0.8, edgecolor='none'))
        if i == num_players - 1:
            ax.set_xlabel("Match Time (seconds)", fontsize=12, fontweight='bold')

    fig.suptitle("Sprint Analysis: Speed Profiles (Sorted by Speed)", fontsize=16, y=1.01)
    plt.tight_layout()
    plt.savefig("output_videos/Chart_1_Speed_Subplots.png", bbox_inches='tight')
    plt.close()

def plot_individual_radars(nts_df):
    top_df = nts_df.head(5).copy()
    if top_df.empty: return
    if 'Threat Score' not in top_df.columns:
        top_df['Threat Score'] = top_df['Grit Score'] * 0.5 + top_df['Influence Score'] * 0.5 

    categories = ['Grit Score', 'Influence Score', 'Leadership Rating', 'Threat Score']
    N = len(categories)
    angles = [n / float(N) * 2 * pi for n in range(N)]
    angles += angles[:1]
    
    fig, axes = plt.subplots(5, 1, figsize=(6, 25), subplot_kw={'projection': 'polar'})
    colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd']
    
    num_plots = min(len(top_df), 5)
    for i in range(5):
        if i >= num_plots:
            axes[i].axis('off')
            continue
        ax = axes[i]
        row = top_df.iloc[i]
        values = [row['Grit Score'], row['Influence Score'], row['Leadership Rating'], row['Threat Score']]
        values += values[:1] 
        color = colors[i % len(colors)]
        ax.plot(angles, values, linewidth=3, linestyle='solid', color=color)
        ax.fill(angles, values, color=color, alpha=0.3)
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(['Grit', 'Influence', 'Lead.', 'Threat'], fontsize=11, fontweight='bold')
        for j, angle in enumerate(angles[:-1]):
            val = values[j]
            ax.text(angle, val + 25, f"{int(val)}", color=color, fontsize=12, fontweight='bold', ha='center', va='center')
        ax.set_yticklabels([]) 
        ax.set_ylim(0, 130)    
        ax.set_title(f"#{int(row['Player ID'])}: {row['Archetype']}", size=16, weight='bold', color=color, pad=20)

    plt.suptitle("NTS Scouting Profiles: Top 5 Players", fontsize=20, fontweight='bold', y=0.99)
    plt.tight_layout()
    plt.savefig("output_videos/Chart_2_Individual_Radars.png", bbox_inches='tight', dpi=150)
    plt.close()

def plot_scatter_quadrants(df):
    print("Generating Physical Profile Scatter Plot...")
    sns.set_style("whitegrid")
    fig, ax = plt.subplots(figsize=(12, 10))
    x = df['Total Distance (m)']
    y = df['Max Speed (km/h)']
    
    avg_dist = x.mean() if not x.empty else 0
    avg_speed = y.mean() if not y.empty else 0
    x_max = x.max() * 1.1 if x.max() > 0 else 100
    y_max = y.max() * 1.1 if y.max() > 0 else 35
    
    ax.fill_between([avg_dist, x_max], avg_speed, y_max, color='green', alpha=0.05)
    ax.fill_between([avg_dist, x_max], 0, avg_speed, color='blue', alpha=0.05)
    ax.fill_between([0, avg_dist], avg_speed, y_max, color='orange', alpha=0.05)
    ax.fill_between([0, avg_dist], 0, avg_speed, color='gray', alpha=0.05)

    palette = {'Team Leader': '#d62728', 'Playmaker': '#9467bd', 'Workhorse': '#2ca02c', 'Speedster': '#ff7f0e', 'Squad Player': '#7f7f7f'}
    sns.scatterplot(x=x, y=y, hue=df['Archetype'], palette=palette, s=300, edgecolor='black', linewidth=1.5, alpha=0.9, ax=ax, zorder=3)
    
    ax.axvline(avg_dist, color='black', linestyle='--', linewidth=2, alpha=0.5)
    ax.axhline(avg_speed, color='black', linestyle='--', linewidth=2, alpha=0.5)
    
    for i in range(len(df)):
        txt = str(int(df.iloc[i]['Player ID']))
        ax.annotate(txt, (x.iloc[i], y.iloc[i]), xytext=(0, 0), textcoords='offset points', ha='center', va='center', fontsize=11, fontweight='bold', color='white')

    ax.text(x_max*0.95, y_max*0.95, "ELITE ATHLETE", ha='right', va='top', fontsize=12, fontweight='bold', color='green', alpha=0.6)
    ax.set_xlabel('Total Distance Covered (meters)', fontsize=14, fontweight='bold')
    ax.set_ylabel('Top Speed Reached (km/h)', fontsize=14, fontweight='bold')
    ax.set_title('Physical Profile Analysis', fontsize=20, fontweight='bold', pad=20)
    ax.set_xlim(0, x_max)
    ax.set_ylim(0, y_max)
    plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize=12, title="Player Archetype")
    plt.tight_layout()
    plt.savefig("output_videos/Chart_3_Scatter_Roles.png", dpi=150, bbox_inches='tight')
    plt.close()

# =================================================================
#  MAIN EXECUTION
# =================================================================

def main():
    # 1. READ & TRACK
    print("Reading video...")
    video_frames = read_video('input_videos/08fd33_4.mp4')
    tracker = Tracker('models/best.pt')
    
    tracks = tracker.get_object_tracks(video_frames, read_from_stub=True, stub_path='stubs/track_stubs.pkl')
    tracks["ball"] = tracker.interpolate_ball_positions(tracks["ball"])
    tracker.add_position_to_tracks(tracks)

    # 2. TRANSFORM & ESTIMATE
    print("Estimating camera movement and speeds...")
    camera_movement_estimator = CameraMovementEstimator(video_frames[0])
    camera_movement_per_frame = camera_movement_estimator.get_camera_movement(video_frames, read_from_stub=True, stub_path='stubs/camera_movement_stub.pkl')
    camera_movement_estimator.add_adjust_positions_to_tracks(tracks, camera_movement_per_frame)

    view_transformer = ViewTransformer()
    view_transformer.add_transformed_position_to_tracks(tracks)

    speed_and_distance_estimator = SpeedAndDistance_Estimator()
    speed_and_distance_estimator.add_speed_and_distance_to_tracks(tracks)
    smooth_speed_data(tracks, window_size=9) 

    # 3. ASSIGN TEAMS & BALL (With STICKY POSSESSION)
    print("Assigning Teams & Possession...")
    team_assigner = TeamAssigner()
    team_assigner.assign_team_color(video_frames[0], tracks['players'][0])
    
    color_1_sum = np.sum(team_assigner.team_colors[1])
    color_2_sum = np.sum(team_assigner.team_colors[2])
    white_team_id = 1 if color_1_sum > color_2_sum else 2

    # A. Team Assignment
    for frame_num, player_track in enumerate(tracks['players']):
        for player_id, track in player_track.items():
            team = team_assigner.get_player_team(video_frames[frame_num], track['bbox'], player_id)
            tracks['players'][frame_num][player_id]['team'] = team 
            tracks['players'][frame_num][player_id]['team_color'] = team_assigner.team_colors[team] if team != 99 else (0,0,0)

    # B. Ball Assignment with MEMORY (Sticky Logic)
    player_assigner = PlayerBallAssigner()
    team_ball_control = []
    
    last_known_carrier = -1 # Memory variable

    for frame_num, player_track in enumerate(tracks['players']):
        ball_bbox = tracks['ball'][frame_num][1]['bbox']
        
        # New Greedy Assignment (finds closest player)
        assigned_player = player_assigner.assign_ball_to_player(player_track, ball_bbox)

        if assigned_player != -1:
            # We found a carrier this frame
            last_known_carrier = assigned_player
            tracks['players'][frame_num][assigned_player]['has_ball'] = True
            team_ball_control.append(tracks['players'][frame_num][assigned_player]['team'])
        
        else:
            # No assignment found. Check Memory.
            # If the last known carrier is still in the video frame, we assume they still have it
            if last_known_carrier != -1 and last_known_carrier in tracks['players'][frame_num]:
                tracks['players'][frame_num][last_known_carrier]['has_ball'] = True
                team_ball_control.append(tracks['players'][frame_num][last_known_carrier]['team'])
            else:
                # No memory, no ball -> No control
                team_ball_control.append(team_ball_control[-1] if team_ball_control else -1)

    team_ball_control = np.array(team_ball_control)

    # 4. STATS PREP
    print("Filtering data & Initializing Stats...")
    drawing_tracks = copy.deepcopy(tracks)
    match_stats = {} 

    for frame_num, player_track in enumerate(drawing_tracks['players']):
        # Filter stats for non-white team (optional, based on your preference)
        ids_to_remove = [pid for pid, trk in player_track.items() if trk.get('team') != white_team_id]
        for pid in ids_to_remove:
            del drawing_tracks['players'][frame_num][pid]

        for player_id, track in drawing_tracks['players'][frame_num].items():
            if 'speed' not in track: track['speed'] = 0
            if 'distance' not in track: track['distance'] = 0

            if player_id not in match_stats:
                match_stats[player_id] = {
                    "Player ID": player_id,
                    "Total Distance (m)": 0, "Max Speed (km/h)": 0,
                    "Frames with Ball": 0, "Pressing Count": 0, "Holding Count": 0, "Idle Count": 0
                }

    # 5. GENERATE VIDEO (AR OVERLAY + SHADOWS)
    print("Generating AR Analysis Video...")
    behavior_analyzer = BehaviorAnalyzer()
    
    # NOTE: We use 'drawing_tracks' (filtered) for the text overlay to avoid clutter
    output_video_frames = tracker.draw_annotations(video_frames, drawing_tracks, team_ball_control)
    output_video_frames = camera_movement_estimator.draw_camera_movement(output_video_frames, camera_movement_per_frame)
    output_video_frames = speed_and_distance_estimator.draw_speed_and_distance(output_video_frames, drawing_tracks)
    
    for frame_num, frame in enumerate(output_video_frames):
        if frame_num >= len(tracks['players']): break
        
        current_frame_tracks = tracks['players'][frame_num] # Use ALL players for AR lines/Shadows
        drawing_frame_tracks = drawing_tracks['players'][frame_num] # Use filtered players for stats

        ball_data = tracks['ball'][frame_num]
        ball_pos_meters = None
        if 1 in ball_data:
            ball_pos_meters = ball_data[1].get('position_transformed')
        
        current_team_control = team_ball_control[frame_num]
        
        # A. STATS & BEHAVIOR
        for player_id, player_data in drawing_frame_tracks.items():
            behavior = behavior_analyzer.analyze_behavior(
                player_id, player_data, current_frame_tracks, ball_pos_meters, current_team_control
            )
            if not behavior: behavior = {"action": "IDLE", "color": (220, 220, 220)}
            drawing_tracks['players'][frame_num][player_id]['behavior'] = behavior

            stats = match_stats[player_id]
            stats["Total Distance (m)"] = max(stats["Total Distance (m)"], player_data.get('distance', 0))
            if player_data.get('speed', 0) > stats["Max Speed (km/h)"]:
                stats["Max Speed (km/h)"] = player_data.get('speed', 0)
            
            action = behavior['action']
            if action == "PRESSING": stats["Pressing Count"] += 1
            elif "HOLDING" in action: stats["Holding Count"] += 1
            elif action == "IDLE": stats["Idle Count"] += 1
            if player_data.get('has_ball', False): stats["Frames with Ball"] += 1

        output_video_frames[frame_num] = tracker.draw_behavior_annotations(frame, drawing_frame_tracks)

        # B. AR OVERLAY (DIRECT LINES + SHADOWS)
        # We pass 'current_frame_tracks' (all players) so lines can connect to anyone
        output_video_frames[frame_num] = draw_ar_overlay(output_video_frames[frame_num], current_frame_tracks)

    print("Saving AR Video...")
    save_video(output_video_frames, 'output_videos/nts_output_AR.avi')
# ============================================================
    # 6. LLM REPORT GENERATION (WORD DOCUMENT DOSSIER)
    # ============================================================
    print("Calculating Stats and Generating Word Report...")
    
    # --- 1. Calculate Tactical Data (Gray Zones) ---
    total_open_passes = 0
    total_blocked_passes = 0
    
    # Sampling frame loop
    for frame_num in range(0, len(tracks['players']), 10): 
        if frame_num >= len(team_ball_control): break
        
        curr_players = tracks['players'][frame_num]
        ball_carrier = None
        
        # Find carrier
        for pid, pdata in curr_players.items():
            if pdata.get('has_ball', False):
                ball_carrier = pdata
                break
        
        # Check valid positions
        if ball_carrier and ball_carrier.get('position_transformed') is not None:
            carrier_team = ball_carrier['team']
            carrier_pos = ball_carrier['position_transformed']
            
            opponents = []
            teammates = []
            
            for pid, pdata in curr_players.items():
                pos_2d = pdata.get('position_transformed')
                if pos_2d is None: continue 

                if pdata['team'] != carrier_team:
                    opponents.append(pos_2d)
                elif pid != ball_carrier.get('track_id'): 
                    teammates.append(pos_2d)
            
            for mate_pos in teammates:
                dist = np.linalg.norm(np.array(carrier_pos) - np.array(mate_pos))
                if dist < 30: 
                    blocked, _ = is_pass_blocked(carrier_pos, mate_pos, opponents)
                    if blocked: total_blocked_passes += 1
                    else: total_open_passes += 1

    team_tactical_stats = {
        "open_passes": total_open_passes,
        "blocked_passes": total_blocked_passes
    }

    # --- 2. Prepare Stats Data ---
    stats_list = list(match_stats.values())
    if not stats_list:
        print("No stats collected. Exiting.")
        return

    fps = 24
    for stats in stats_list:
        stats["Time with Ball (s)"] = round(stats.get("Frames with Ball", 0) / fps, 2)
        stats["Time Pressing (s)"] = round(stats.get("Pressing Count", 0) / fps, 2)
        stats["Total Distance (m)"] = round(stats.get("Total Distance (m)", 0), 2)
        stats["Max Speed (km/h)"] = round(stats.get("Max Speed (km/h)", 0), 2)
        # Safe cleanup
        for k in ["Frames with Ball", "Pressing Count", "Holding Count", "Idle Count"]:
            if k in stats: del stats[k]

    df = pd.DataFrame(stats_list)
    nts_df = df.copy()
    
    # Normalize
    cols = ['Total Distance (m)', 'Max Speed (km/h)', 'Time with Ball (s)', 'Time Pressing (s)']
    for col in cols:
        if col in nts_df.columns:
            min_val, max_val = nts_df[col].min(), nts_df[col].max()
            if max_val == min_val: nts_df[f'norm_{col}'] = 0
            else: nts_df[f'norm_{col}'] = (nts_df[col] - min_val) / (max_val - min_val)
        else: nts_df[f'norm_{col}'] = 0

    nts_df['Grit Score'] = ((nts_df['norm_Total Distance (m)'] * 0.6) + (nts_df['norm_Time Pressing (s)'] * 0.4)) * 100
    nts_df['Influence Score'] = nts_df['norm_Time with Ball (s)'] * 100
    nts_df['Leadership Rating'] = ((nts_df['Grit Score'] * 0.5) + (nts_df['Influence Score'] * 0.5))

    def classify_role(row):
        if row['Leadership Rating'] >= 70: return "Team Leader"
        if row['Grit Score'] >= 75: return "Workhorse"
        if row['Influence Score'] >= 75: return "Playmaker"
        if row['Max Speed (km/h)'] > 30: return "Speedster" 
        return "Squad Player"

    nts_df['Archetype'] = nts_df.apply(classify_role, axis=1)
    
    # --- 3. GENERATE WORD DOC DOSSIER ---
    print("Generating Narrative Dossier (Word Doc)...")
    
    try:
        from narrative_generator import generate_nts_report
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Initialize Word Doc
        doc = Document()
        
        # Title
        title = doc.add_heading('NTS Psychological Profile Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Match Analysis | Automated Computer Vision Assessment")
        doc.add_paragraph("--------------------------------------------------------")

        # Select Top 3 Players
        top_players = nts_df.sort_values(by='Leadership Rating', ascending=False).head(3)
        
        for index, player_row in top_players.iterrows():
            pid = int(player_row['Player ID'])
            role = player_row['Archetype']
            print(f" > Writing dossier for Player {pid}...")

            # Generate Text
            report_text = generate_nts_report(player_row, team_tactical_stats)
            
            # Add to Word Doc
            # 1. Player Header
            heading = doc.add_heading(f"Player #{pid} - {role}", level=1)
            
            # 2. The AI Narrative
            p = doc.add_paragraph(report_text)
            p_format = p.paragraph_format
            p_format.space_after = Pt(12)
            
            # 3. Supporting Evidence (Bullet points)
            doc.add_heading('Behavioral Evidence:', level=3)
            
            # Formatting bullet points
            bullets = [
                f"Cognitive Load: Operated in high-pressure zones ({team_tactical_stats['blocked_passes']} blocked lanes).",
                f"Composure: Retained possession for {player_row['Time with Ball (s)']}s.",
                f"Work Ethic: Covered {player_row['Total Distance (m)']}m with {player_row['Time Pressing (s)']}s of pressing."
            ]
            for b in bullets:
                doc.add_paragraph(b, style='List Bullet')
            
            doc.add_paragraph("_" * 50) # Separator line

        # Save the file
        output_filename = "output_videos/NTS_Psych_Profile.docx"
        doc.save(output_filename)
        print(f"Report saved successfully to: {output_filename}")

        # Generate Plots (Standard)
        top_ids_for_charts = top_players['Player ID'].tolist()
        plot_speed_subplots(drawing_tracks, top_ids_for_charts, stats_list)
        plot_individual_radars(top_players)
        plot_scatter_quadrants(nts_df)

    except ImportError as e:
        print(f"Missing library: {e}. Run 'pip install python-docx'")
    except Exception as e:
        print(f"Error generating report: {e}")

    print("Analysis Complete!")

if __name__ == '__main__':
    main()